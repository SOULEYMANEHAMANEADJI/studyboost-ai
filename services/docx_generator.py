"""Générateur DOCX (Word) à partir de Markdown pour StudyBoost AI."""
from __future__ import annotations

import re
from datetime import datetime
from io import BytesIO

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


def markdown_to_docx(text: str, title: str | None = None) -> bytes:
    if not title:
        title = f"StudyBoost_{datetime.now().strftime('%Y%m%d_%H%M')}"

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.15

    if "Heading 1" in [s.name for s in doc.styles]:
        h1 = doc.styles["Heading 1"]
        h1.font.size = Pt(20)
        h1.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
        h1.font.bold = True
    if "Heading 2" in [s.name for s in doc.styles]:
        h2 = doc.styles["Heading 2"]
        h2.font.size = Pt(15)
        h2.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
        h2.font.bold = True
    if "Heading 3" in [s.name for s in doc.styles]:
        h3 = doc.styles["Heading 3"]
        h3.font.size = Pt(12)
        h3.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
        h3.font.bold = True

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(title)
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
    run.bold = True

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(f"Généré par StudyBoost AI — {datetime.now().strftime('%d/%m/%Y')}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    run.italic = True

    doc.add_paragraph()

    lines = text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("> "):
            p = doc.add_paragraph(line[2:])
            p.paragraph_format.left_indent = Cm(1)
            p.runs[0].italic = True
            p.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        elif line.lstrip().startswith(("- ", "* ")):
            while i < len(lines) and lines[i].lstrip().startswith(("- ", "* ")):
                item = lines[i].lstrip()[2:]
                p = doc.add_paragraph(item, style="List Bullet")
                _apply_inline_format(p, item)
                i += 1
            continue
        elif re.match(r"^\d+\.\s", line):
            while i < len(lines) and re.match(r"^\d+\.\s", lines[i]):
                item = re.sub(r"^\d+\.\s", "", lines[i])
                p = doc.add_paragraph(item, style="List Number")
                _apply_inline_format(p, item)
                i += 1
            continue
        elif line in ("---", "***", "___"):
            doc.add_paragraph("─" * 60)
        else:
            p = doc.add_paragraph()
            _apply_inline_format(p, line)

        i += 1

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _apply_inline_format(paragraph, text: str):
    parts = re.split(r"(\*\*\*.+?\*\*\*|\*\*.+?\*\*|\*.+?\*|`.+?`)", text)
    for part in parts:
        if part.startswith("***") and part.endswith("***"):
            run = paragraph.add_run(part[3:-3])
            run.bold = True
            run.italic = True
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x7C, 0x3A, 0xED)
        elif part.strip():
            run = paragraph.add_run(part)
